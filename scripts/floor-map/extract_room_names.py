#!/usr/bin/env python3
"""
Извлекает официальный список помещений из Word-файла (.docx) и сохраняет его
в data/floor-map/room-names.json.

Скрипт НИЧЕГО не додумывает: если номер или название отсутствуют, запись всё
равно попадает в JSON с пометкой в поле "issues". Этаж определяется только по
первой цифре номера комнаты (101 -> 1, 201 -> 2, 301 -> 3); для помещений без
номера этаж остаётся null и помечается как floor_unknown.

Использование:
    pip install python-docx
    python3 scripts/floor-map/extract_room_names.py <путь-к-.docx> [выходной-json]
"""

import json
import re
import sys
from pathlib import Path

try:
    import docx  # python-docx
except ImportError:
    sys.exit("Нужен python-docx: pip install python-docx")

# Строки-заголовки внутри таблицы («1 этаж», «Список» и т. п.)
HEADER_RE = re.compile(r"^\s*(\d+\s*этаж|подвал|цоколь|№)\s*$", re.IGNORECASE)
NUMBER_RE = re.compile(r"^\s*(\d{3})\s*([а-яa-z])?\s*$", re.IGNORECASE)


def dedupe(cells):
    """Схлопывает объединённые ячейки (python-docx повторяет их текст)."""
    out = []
    for c in cells:
        if not out or out[-1] != c:
            out.append(c)
    return out


def parse(docx_path: Path):
    document = docx.Document(str(docx_path))
    if not document.tables:
        sys.exit(f"В {docx_path} нет таблиц — нечего разбирать.")

    rooms = []
    current_section = None

    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            cells = dedupe([c.text.strip() for c in row.cells])
            cells = [c for c in cells if c != ""]
            if not cells:
                continue

            # Заголовок секции («1 этаж») — одна ячейка на всю строку
            if len(cells) == 1 and HEADER_RE.match(cells[0]):
                current_section = cells[0]
                continue
            if cells[0] == "№":  # шапка таблицы
                continue

            issues = []
            match = NUMBER_RE.match(cells[0])
            if match:
                number = match.group(1) + (" " + match.group(2) if match.group(2) else "")
                name = cells[1] if len(cells) > 1 else ""
                floor = int(match.group(1)[0])
                floor_source = "number_prefix"
            else:
                # Строка без номера: в первой ячейке сразу название
                number = None
                name = cells[0]
                floor = None
                floor_source = None
                issues.append("missing_number")
                issues.append("floor_unknown")

            if not name:
                issues.append("missing_name")

            rooms.append(
                {
                    "id": f"room-{number.replace(' ', '')}" if number else None,
                    "number": number,
                    "name": name or None,
                    "floor": floor,
                    "floor_source": floor_source,
                    "section_header": current_section,
                    "source_table": table_index,
                    "source_row": row_index,
                    "issues": issues,
                }
            )

    return rooms


def main():
    if len(sys.argv) < 2:
        sys.exit(__doc__)
    src = Path(sys.argv[1])
    dst = Path(sys.argv[2]) if len(sys.argv) > 2 else Path("data/floor-map/room-names.json")

    rooms = parse(src)
    dst.parent.mkdir(parents=True, exist_ok=True)
    payload = {
        "source_file": src.name,
        "total": len(rooms),
        "rooms": rooms,
    }
    dst.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    by_floor = {}
    for r in rooms:
        by_floor[r["floor"]] = by_floor.get(r["floor"], 0) + 1
    print(f"Записано {len(rooms)} помещений в {dst}")
    for floor in sorted(by_floor, key=lambda f: (f is None, f)):
        label = f"этаж {floor}" if floor is not None else "этаж не определён"
        print(f"  {label}: {by_floor[floor]}")
    problems = [r for r in rooms if r["issues"]]
    print(f"Записей с вопросами: {len(problems)}")
    for r in problems:
        print(f"  строка {r['source_row']}: {r['number'] or '—'} / {r['name'] or '—'} -> {', '.join(r['issues'])}")


if __name__ == "__main__":
    main()
